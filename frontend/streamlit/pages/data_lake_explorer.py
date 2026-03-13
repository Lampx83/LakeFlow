# frontend/streamlit/pages/data_lake_explorer.py

from pathlib import Path
import hashlib
import json
import os
import sqlite3

import pandas as pd
import streamlit as st

from config.settings import DATA_ROOT
from state.session import require_login
from services.api_client import get_data_path_from_api
from utils.sqlite_viewer import copy_db_to_temp

# File viewer: size limits (avoid hanging)
MAX_VIEW_TEXT_BYTES = 10 * 1024 * 1024   # 10 MB for txt/json/jsonl
MAX_VIEW_NPY_BYTES = 5 * 1024 * 1024     # 5 MB for npy
MAX_VIEW_PDF_BYTES = 50 * 1024 * 1024    # 50 MB for pdf
MAX_JSONL_LINES = 500

# =====================================================
# DATA ZONES
# =====================================================

ZONE_NAMES = [
    "000_inbox",
    "100_raw",
    "200_staging",
    "300_processed",
    "400_embeddings",
    "500_catalog",
]

def _zones_from_root(root: Path) -> dict[str, Path]:
    return {z: root / z for z in ZONE_NAMES}

# Fallback when API not yet called (module load)
ZONES = _zones_from_root(DATA_ROOT)

MAX_TREE_DEPTH = 30  # Limit depth to avoid infinite recursion
CACHE_TTL_TREE = 90  # Seconds to cache list_dir (NAS slow)

# Pipeline steps: 0=inbox→raw, 1=raw→staging, 2=staging→processed, 3=processed→embeddings, 4=embeddings→Qdrant
PIPELINE_STEP_LABELS = {
    0: "Step 0 (Inbox → Raw)",
    1: "Step 1 (Raw → Staging)",
    2: "Step 2 (Staging → Processed)",
    3: "Step 3 (Processed → Embeddings)",
    4: "Step 4 (Embeddings → Qdrant)",
}

# Step status columns for files in 000_inbox
INBOX_STEP_COLUMNS = ["Ingest", "Staging", "Processed", "Embeddings", "Qdrant"]


# =====================================================
# NAS READ CACHE (reduce lag, avoid re-reading same path)
# =====================================================

@st.cache_data(ttl=CACHE_TTL_TREE)
def _list_dir_cached(path_str: str) -> tuple[list[str], list[tuple[str, int]]]:
    """
    Read directory from NAS, return (dir names, [(file name, size)]).
    Uses thread to avoid blocking UI; result is cached.
    """
    path = Path(path_str)
    dirs, files = [], []
    try:
        for p in sorted(path.iterdir()):
            if p.name.startswith("."):
                continue
            if p.is_dir():
                dirs.append(p.name)
            elif p.is_file():
                try:
                    files.append((p.name, p.stat().st_size))
                except OSError:
                    files.append((p.name, 0))
    except (PermissionError, OSError):
        pass
    return (sorted(dirs, key=str.lower), sorted(files, key=lambda x: x[0].lower()))


# =====================================================
# TREE VIEW — LAZY LOADING
# =====================================================

def _format_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 * 1024):.1f} MB"


def _path_to_key(path_str: str) -> str:
    """Create unique widget key from path (avoid collisions when path is long)."""
    return hashlib.md5(path_str.encode()).hexdigest()[:24]


@st.cache_data(ttl=120)
def _sha256_file_cached(path_str: str) -> str | None:
    """Compute SHA256 of file (to compare with raw_objects). Cache 120s."""
    path = Path(path_str)
    if not path.is_file():
        return None
    try:
        h = hashlib.sha256()
        with path.open("rb") as f:
            while chunk := f.read(1024 * 1024):
                h.update(chunk)
        return h.hexdigest()
    except (OSError, PermissionError):
        return None


def get_inbox_file_pipeline_steps(file_path: Path, domain: str) -> dict[str, str]:
    """
    For file in 000_inbox: return whether each step has been processed.
    Keys: Ingest, Staging, Processed, Embeddings, Qdrant. Value: "✓" or "" (Qdrant may be "?" if unknown).
    """
    result = {k: "" for k in INBOX_STEP_COLUMNS}
    path_str = str(file_path.resolve())
    file_hash = _sha256_file_cached(path_str)
    if not file_hash:
        return result
    root = DATA_ROOT
    catalog_db = root / "500_catalog" / "catalog.sqlite"
    # Step 0: Ingest (in raw_objects) — copy DB to temp to avoid Errno 35
    temp_path = None
    try:
        if catalog_db.exists():
            temp_path = copy_db_to_temp(catalog_db)
            conn = sqlite3.connect(str(temp_path), timeout=5)
            cur = conn.execute("SELECT 1 FROM raw_objects WHERE hash = ? LIMIT 1", (file_hash,))
            if cur.fetchone():
                result["Ingest"] = "✓"
            conn.close()
    except Exception:
        pass
    finally:
        if temp_path and Path(temp_path).exists():
            try:
                Path(temp_path).unlink()
            except OSError:
                pass
    if result["Ingest"] != "✓":
        return result
    # Steps 1–3: check directories; support both domain/hash and hash (legacy)
    _staging_dir = root / "200_staging" / domain / file_hash if domain and domain != "." else root / "200_staging" / file_hash
    _staging_alt = root / "200_staging" / file_hash if domain and domain != "." else None
    if (_staging_dir / "validation.json").exists() or (_staging_alt and (_staging_alt / "validation.json").exists()):
        result["Staging"] = "✓"
    _processed_dir = root / "300_processed" / domain / file_hash if domain and domain != "." else root / "300_processed" / file_hash
    _processed_alt = root / "300_processed" / file_hash if domain and domain != "." else None
    if (_processed_dir / "chunks.json").exists() or (_processed_alt and (_processed_alt / "chunks.json").exists()):
        result["Processed"] = "✓"
    _emb_dir = root / "400_embeddings" / domain / file_hash if domain and domain != "." else root / "400_embeddings" / file_hash
    _emb_alt = root / "400_embeddings" / file_hash if domain and domain != "." else None
    if (_emb_dir / "embedding.npy").exists() or (_emb_alt and (_emb_alt / "embedding.npy").exists()):
        result["Embeddings"] = "✓"
    # Step 4: Qdrant — not in catalog, leave blank or "?"
    result["Qdrant"] = "?" if result["Embeddings"] == "✓" else ""
    return result


def get_raw_file_pipeline_steps(file_path: Path, domain: str) -> dict[str, str]:
    """
    For file in 100_raw: Ingest always ✓ (already in Raw), later steps check by domain/hash.
    file_hash = file_path.stem (file name without extension).
    """
    result = {k: "" for k in INBOX_STEP_COLUMNS}
    result["Ingest"] = "✓"  # Already ingested (file is in Raw)
    file_hash = file_path.stem
    domain = domain or "."
    root = DATA_ROOT
    # Support both domain/hash and hash (legacy); backend writes embedding.npy (no 's')
    _staging = root / "200_staging" / domain / file_hash if domain != "." else root / "200_staging" / file_hash
    _staging_alt = root / "200_staging" / file_hash if domain != "." else None
    if (_staging / "validation.json").exists() or (_staging_alt and (_staging_alt / "validation.json").exists()):
        result["Staging"] = "✓"
    _processed = root / "300_processed" / domain / file_hash if domain != "." else root / "300_processed" / file_hash
    _processed_alt = root / "300_processed" / file_hash if domain != "." else None
    if (_processed / "chunks.json").exists() or (_processed_alt and (_processed_alt / "chunks.json").exists()):
        result["Processed"] = "✓"
    _emb = root / "400_embeddings" / domain / file_hash if domain != "." else root / "400_embeddings" / file_hash
    _emb_alt = root / "400_embeddings" / file_hash if domain != "." else None
    if (_emb / "embedding.npy").exists() or (_emb_alt and (_emb_alt / "embedding.npy").exists()):
        result["Embeddings"] = "✓"
    result["Qdrant"] = "?" if result["Embeddings"] == "✓" else ""
    return result


def render_folder_tree(
    root: Path,
    zone_name: str,
    expanded_set: set[str],
    current_folder: str | None,
    zone_root: Path,
    depth: int = 0,
) -> None:
    """
    Show directory tree only (no files). Click folder = expand/collapse + select to show file list on the right.
    """
    if depth >= MAX_TREE_DEPTH:
        st.caption("… (depth limit reached)")
        return

    path_str = str(root.resolve())
    try:
        dir_names, _ = _list_dir_cached(path_str)
    except Exception as e:
        st.caption(f"⚠️ Error: {e}")
        return

    indent = "  " * depth  # 2 spaces for compact view

    for d_name in dir_names:
        child_path = root / d_name
        child_str = str(child_path.resolve())
        key_suffix = _path_to_key(child_str)
        is_selected = current_folder == child_str
        icon = "▼" if child_str in expanded_set else "▶"
        try:
            child_dirs, child_files = _list_dir_cached(child_str)
            count = len(child_dirs) + len(child_files)
            count_str = f" ({count})"
        except Exception:
            count_str = ""
        label = f"{indent}{icon} {d_name}{count_str}" + (" ✓" if is_selected else "")

        if st.button(label, key=f"tree_{zone_name}_{key_suffix}", type="primary" if is_selected else "secondary"):
            if child_str in expanded_set:
                expanded_set.discard(child_str)
            else:
                expanded_set.add(child_str)
            st.session_state["datalake_current_folder"] = child_str
            st.rerun()

        if child_str in expanded_set:
            render_folder_tree(child_path, zone_name, expanded_set, current_folder, zone_root, depth + 1)


def render_file_list(
    folder_path: Path,
    zone_name: str,
    zone_root: Path,
    selected_file: str | None,
) -> None:
    """
    Show file list in directory as table; select row to view content.
    """
    if not _is_safe_path(folder_path, zone_root):
        st.warning("Directory is not in zone.")
        return
    try:
        _, file_infos = _list_dir_cached(str(folder_path.resolve()))
    except Exception as e:
        st.warning(f"Cannot read directory: {e}")
        return

    if not file_infos:
        st.info("Directory is empty or has no files.")
        return

    # 000_inbox or 100_raw (root or domain dir): table has step columns with ✓
    # 100_raw skips Ingest (always ✓ because file is in Raw)
    # Other zones: table has Pipeline step (text description)
    rows = []
    is_inbox_zone = zone_name == "000_inbox"
    is_raw_zone = zone_name == "100_raw"
    use_step_columns = is_inbox_zone or is_raw_zone
    # Domain = first segment under zone (correct even when viewing file in subdir, e.g. 000_inbox/education/2024/)
    try:
        rel = folder_path.resolve().relative_to(zone_root.resolve())
        domain = rel.parts[0] if rel.parts else "."
    except (ValueError, OSError):
        domain = "." if folder_path.resolve() == zone_root.resolve() else folder_path.name
    for f_name, size in file_infos:
        file_path = folder_path / f_name
        _path = str(file_path.resolve())
        if use_step_columns and is_inbox_zone:
            steps = get_inbox_file_pipeline_steps(file_path, domain)
            rows.append({
                "File name": f_name,
                "Size": _format_size(size),
                **{col: steps.get(col, "") for col in INBOX_STEP_COLUMNS},
                "_path": _path,
            })
        elif use_step_columns and is_raw_zone:
            steps = get_raw_file_pipeline_steps(file_path, domain)
            rows.append({
                "File name": f_name,
                "Size": _format_size(size),
                **{col: steps.get(col, "") for col in INBOX_STEP_COLUMNS},
                "_path": _path,
            })
        else:
            step_label = get_pipeline_step_for_path(
                file_path, zone_name, data_root=zone_root.parent
            )
            step_str = step_label if step_label else "—"
            rows.append({
                "File name": f_name,
                "Size": _format_size(size),
                "Pipeline step": step_str,
                "_path": _path,
            })

    if use_step_columns:
        df = pd.DataFrame([{k: r[k] for k in ["File name", "Size"] + INBOX_STEP_COLUMNS} for r in rows])
        if is_inbox_zone:
            st.caption("✓ = step processed. Ingest = Step 0 (→ Raw), Staging = Step 1, Processed = Step 2, Embeddings = Step 3, Qdrant = Step 4 (?) if not confirmed.")
        else:
            st.caption("✓ = step processed. Files in Raw so Ingest always ✓. Staging = Step 1, Processed = Step 2, Embeddings = Step 3, Qdrant = Step 4 (?) if not confirmed.")
    else:
        df = pd.DataFrame([{"File name": r["File name"], "Size": r["Size"], "Pipeline step": r["Pipeline step"]} for r in rows])

    st.markdown(f"**Files in** `{folder_path.name}`")
    st.dataframe(df, use_container_width=True, hide_index=True)

    chosen = st.selectbox(
        "Select file to view content",
        options=[r["_path"] for r in rows],
        format_func=lambda p: Path(p).name,
        index=next((i for i, r in enumerate(rows) if r["_path"] == selected_file), 0),
        key=f"file_sel_{zone_name}_{_path_to_key(str(folder_path.resolve()))}",
    )
    if chosen:
        st.session_state["datalake_selected_file"] = chosen


def get_pipeline_step_for_path(
    file_path: Path, zone_name: str, data_root: Path | None = None
) -> str | None:
    """
    Determine which pipeline step (0–4) the file has been processed to, using catalog and filesystem checks.
    Returns step description or None if unknown.
    data_root: if set, use it (path from backend); else use DATA_ROOT from env.
    """
    root = data_root if data_root is not None else DATA_ROOT
    staging_root = root / "200_staging"
    processed_root = root / "300_processed"
    embeddings_root = root / "400_embeddings"
    catalog_db = root / "500_catalog" / "catalog.sqlite"

    if zone_name == "000_inbox":
        return "Step 0 not run (file still in Inbox)"
    if zone_name not in ("100_raw", "200_staging", "300_processed", "400_embeddings"):
        return None

    file_hash = None
    domain = None
    zone_root = root / zone_name if data_root is not None else ZONES.get(zone_name)
    if not zone_root:
        return None
    try:
        rel = file_path.resolve().relative_to(zone_root.resolve())
        parts = rel.parts
        if zone_name == "100_raw" and file_path.is_file():
            file_hash = file_path.stem
            domain = parts[0] if len(parts) >= 2 else None
        elif zone_name in ("200_staging", "300_processed", "400_embeddings"):
            if file_path.is_dir():
                file_hash = file_path.name
                domain = parts[0] if len(parts) >= 2 else None
            else:
                file_hash = file_path.parent.name
                domain = parts[0] if len(parts) >= 2 else None
    except ValueError:
        return None
    if not file_hash:
        return None

    domain = domain or "."
    step = -1
    if not catalog_db.exists():
        return "Catalog not ready (Step 0 not run)"

    temp_path = None
    try:
        temp_path = copy_db_to_temp(catalog_db)
        conn = sqlite3.connect(str(temp_path), timeout=5)
        cur = conn.execute("SELECT 1 FROM raw_objects WHERE hash = ? LIMIT 1", (file_hash,))
        if cur.fetchone():
            step = 0
        conn.close()
    except Exception:
        return "Cannot read Catalog"
    finally:
        if temp_path and Path(temp_path).exists():
            try:
                Path(temp_path).unlink()
            except OSError:
                pass

    if step < 0:
        return "Not in Catalog (Step 0 not run)"

    domain = domain or "."
    # Support both domain/hash and hash (legacy); backend writes embedding.npy (no 's')
    _staging = staging_root / domain / file_hash if domain != "." else staging_root / file_hash
    _staging_alt = staging_root / file_hash if domain != "." else None
    if (_staging / "validation.json").exists() or (_staging_alt and (_staging_alt / "validation.json").exists()):
        step = 1
    _processed = processed_root / domain / file_hash if domain != "." else processed_root / file_hash
    _processed_alt = processed_root / file_hash if domain != "." else None
    if (_processed / "chunks.json").exists() or (_processed_alt and (_processed_alt / "chunks.json").exists()):
        step = 2
    _emb = embeddings_root / domain / file_hash if domain != "." else embeddings_root / file_hash
    _emb_alt = embeddings_root / file_hash if domain != "." else None
    if (_emb / "embedding.npy").exists() or (_emb_alt and (_emb_alt / "embedding.npy").exists()):
        step = 3
    # Step 4: need to query Qdrant to confirm — not in catalog

    return PIPELINE_STEP_LABELS.get(step, f"Step {step}")


def _is_safe_path(file_path: Path, zone_root: Path) -> bool:
    """Ensure file is within zone (prevent path traversal)."""
    try:
        return file_path.resolve().is_relative_to(zone_root.resolve())
    except (ValueError, OSError):
        return False


def render_file_content(file_path: Path) -> None:
    """
    Display file content by format: txt, json, jsonl, npy, pdf, csv.
    Size limits to avoid hanging.
    """
    if not file_path.is_file():
        st.warning("File does not exist or cannot be read.")
        return

    try:
        size = file_path.stat().st_size
    except OSError:
        st.error("Cannot read file info.")
        return

    suffix = file_path.suffix.lower()

    # ---------- TXT ----------
    if suffix == ".txt":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File too large ({size / (1024*1024):.1f} MB). Only supports viewing file ≤ {MAX_VIEW_TEXT_BYTES // (1024*1024)} MB.")
            _download_button(file_path)
            return
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            st.code(text, language="text")
        except Exception as e:
            st.error(f"Error reading file: {e}")

    # ---------- JSON ----------
    elif suffix == ".json":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File too large. Only supports viewing file ≤ {MAX_VIEW_TEXT_BYTES // (1024*1024)} MB.")
            _download_button(file_path)
            return
        try:
            with file_path.open("r", encoding="utf-8") as f:
                data = json.load(f)
            st.json(data)
        except Exception as e:
            st.error(f"Error reading JSON: {e}")

    # ---------- JSONL ----------
    elif suffix == ".jsonl":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File too large. Only showing max {MAX_JSONL_LINES} lines.")
        try:
            lines = []
            with file_path.open("r", encoding="utf-8", errors="replace") as f:
                for i, line in enumerate(f):
                    if i >= MAX_JSONL_LINES:
                        st.caption(f"… Showing first {MAX_JSONL_LINES} lines only. File may have more.")
                        break
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        lines.append(json.loads(line))
                    except json.JSONDecodeError:
                        lines.append({"raw": line})
            if lines:
                st.dataframe(lines, use_container_width=True)
            else:
                st.info("File is empty or has no valid JSON lines.")
        except Exception as e:
            st.error(f"Error reading file: {e}")

    # ---------- NPY ----------
    elif suffix == ".npy":
        if size > MAX_VIEW_NPY_BYTES:
            st.warning(f"File too large ({size / (1024*1024):.1f} MB). Only supports viewing file ≤ {MAX_VIEW_NPY_BYTES // (1024*1024)} MB.")
            _download_button(file_path)
            return
        try:
            import numpy as np
            arr = np.load(file_path, allow_pickle=False)
            st.write("**Shape:**", arr.shape)
            st.write("**Dtype:**", str(arr.dtype))
            if arr.size <= 100:
                st.write("**Data:**")
                st.write(arr)
            else:
                st.write("**Sample (first 100 elements):**")
                st.write(arr.flat[:100])
        except ImportError:
            st.info("Install `numpy` to view .npy files. You can download the file.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Error reading .npy file: {e}")

    # ---------- PDF ----------
    elif suffix == ".pdf":
        if size > MAX_VIEW_PDF_BYTES:
            st.warning(f"File too large. Only supports file info ≤ {MAX_VIEW_PDF_BYTES // (1024*1024)} MB.")
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            n_pages = len(reader.pages)
            st.write(f"**Page count:** {n_pages}")
            _download_button(file_path)
        except ImportError:
            st.info("Install `pypdf` to view PDF info. You can download the file.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Error reading PDF: {e}")
            _download_button(file_path)

    # ---------- CSV ----------
    elif suffix == ".csv":
        if size > MAX_VIEW_TEXT_BYTES:
            st.warning(f"File too large. Only showing partial content.")
        try:
            import pandas as pd
            df = pd.read_csv(file_path, nrows=1000, encoding="utf-8", on_bad_lines="skip")
            st.dataframe(df, use_container_width=True)
            if size > 1024 * 1024:
                st.caption("Showing first 1000 lines only.")
        except ImportError:
            st.info("Install `pandas` to view CSV. You can download the file.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Error reading CSV: {e}")
    # Word/Excel viewer section
    elif suffix == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            # Extract full text from paragraphs
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract from tables (to avoid missing info like application counts)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    full_text.append(f" [Table] {row_text}")
            
            st.write(f"**Format:** Microsoft Word")
            st.text_area("Extracted content:", value="\n".join(full_text), height=400)
            _download_button(file_path)
        except ImportError:
            st.info("Install `python-docx` to view Word content. You can download the file.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Error reading Word: {e}")
            _download_button(file_path)
    elif suffix in [".xls", ".xlsx"]:
        try:
            import pandas as pd
            # Read first Excel sheet
            df = pd.read_excel(file_path)
            
            st.write(f"**Format:** Microsoft Excel")
            st.write(f"**Rows:** {len(df)} | **Columns:** {len(df.columns)}")
            
            # Display data table
            st.dataframe(df, use_container_width=True)
            _download_button(file_path)
        except ImportError:
            st.info("Install `pandas`, `openpyxl` and `xlrd` to view Excel. You can download the file.")
            _download_button(file_path)
        except Exception as e:
            st.error(f"Error reading Excel: {e}")
            _download_button(file_path)
    # ---------- Other formats ----------
    else:
        st.info(f"Format `{suffix}` not supported for direct viewing. You can download the file.")
        _download_button(file_path)


def _download_button(file_path: Path) -> None:
    try:
        data = file_path.read_bytes()
        st.download_button(
            "⬇️ Download file",
            data=data,
            file_name=file_path.name,
            mime="application/octet-stream",
            key=f"dl_{hashlib.md5(str(file_path).encode()).hexdigest()[:16]}",
        )
    except Exception:
        pass


# =====================================================
# MAIN PAGE
# =====================================================

def render():
    if not require_login():
        return

    # Get data root from backend (correct path when running dev; avoid default /data)
    if "data_lake_root" not in st.session_state:
        api_path = get_data_path_from_api()
        if api_path:
            st.session_state["data_lake_root"] = Path(api_path).expanduser().resolve()
        else:
            st.session_state["data_lake_root"] = DATA_ROOT
    effective_root = st.session_state["data_lake_root"]
    zones = _zones_from_root(effective_root)

    st.header("🗂️ Data Lake Explorer")
    st.caption(
        "View Data Lake structure. Click a file to show content and **pipeline step** (from Catalog: Step 0–3 run). "
        "Catalog (500_catalog) contains raw_objects, ingest_log."
    )

    # --------------------------------------------------
    # SELECT ZONE
    # --------------------------------------------------
    zone_name = st.selectbox(
        "📂 Select data zone",
        list(zones.keys()),
    )

    zone_path = zones[zone_name]

    if not zone_path.exists():
        st.warning(f"Zone does not exist: {zone_path}")
        return

    st.subheader(f"📁 {zone_name}")

    if "datalake_expanded" not in st.session_state:
        st.session_state.datalake_expanded = {}
    expanded_set = st.session_state.datalake_expanded.setdefault(zone_name, set())
    zone_root_str = str(zone_path.resolve())
    current_folder = st.session_state.get("datalake_current_folder")
    if not current_folder:
        current_folder = zone_root_str
    else:
        try:
            Path(current_folder).resolve().relative_to(zone_path.resolve())
        except (ValueError, OSError, TypeError):
            current_folder = zone_root_str
    if not current_folder:
        current_folder = zone_root_str

    # --------------------------------------------------
    # 2-COLUMN LAYOUT: Directory tree (compact) left | File table + content right
    # --------------------------------------------------
    col_tree, col_files = st.columns([0.9, 2.1])

    with col_tree:
        st.markdown("**📁 Directory**")
        try:
            root_dirs, root_files = _list_dir_cached(zone_root_str)
            root_count = len(root_dirs) + len(root_files)
            root_label = f"📂 Root ({root_count})"
        except Exception:
            root_label = "📂 Root"
        if st.button(root_label, key="datalake_root", help="View files in zone root directory"):
            st.session_state["datalake_current_folder"] = zone_root_str
            st.rerun()
        st.divider()
        with st.spinner("Loading tree..."):
            render_folder_tree(zone_path, zone_name, expanded_set, current_folder, zone_path)

    with col_files:
        st.markdown("**📄 File list**")
        folder_path = Path(current_folder)

        with st.spinner("Loading list..."):
            render_file_list(folder_path, zone_name, zone_path, st.session_state.get("datalake_selected_file"))

        # File content when selected
        selected = st.session_state.get("datalake_selected_file")
        if selected:
            sel_path = Path(selected)
            if sel_path.is_file() and _is_safe_path(sel_path, zone_path):
                st.divider()
                st.subheader(f"📄 `{sel_path.name}`")
                step_info = get_pipeline_step_for_path(
                    sel_path, zone_name, data_root=effective_root
                )
                if step_info:
                    st.caption(f"🔄 **Pipeline:** {step_info}")
                if st.button("✕ Close file view", key="datalake_close_file"):
                    del st.session_state["datalake_selected_file"]
                    st.rerun()
                render_file_content(sel_path)

    if zone_name == "500_catalog":
        st.caption("View SQLite (catalog) content at **🗄️ SQLite Viewer** on the left sidebar.")

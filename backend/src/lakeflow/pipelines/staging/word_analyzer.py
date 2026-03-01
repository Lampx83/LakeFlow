from pathlib import Path
from typing import Dict, Any
import docx
from docx.opc.exceptions import PackageNotFoundError

class StagingError(RuntimeError):
    """Staging error with clear reason, matches project error handling."""

def analyze_word(file_path: Path) -> Dict[str, Any]:
    """
    Analyze Word file structure in detail (200_staging).
    Includes format validation, metadata extraction, and content assessment.
    """
    if not file_path.exists():
        raise StagingError(f"Word file does not exist: {file_path}")

    try:
        doc = docx.Document(file_path)
    except PackageNotFoundError:
        raise StagingError(f"File is not valid Word (.docx) format or has corrupted ZIP structure.")
    except Exception as e:
        raise StagingError(f"Word analysis error: {str(e)}")

    # 1. Collect metadata (similar to PDF metadata)
    prop = doc.core_properties
    metadata = {
        "author": prop.author,
        "created": str(prop.created) if prop.created else None,
        "last_modified_by": prop.last_modified_by,
        "revision": prop.revision,
        "title": prop.title,
    }

    # 2. Analyze content structure
    paragraphs = doc.paragraphs
    tables = doc.tables
    
    # Check if there is actual text (guard against empty file)
    text_content = [p.text for p in paragraphs if p.text.strip()]
    has_text = len(text_content) > 0
    
    # Compute counts (similar to PDF)
    paragraph_count = len(paragraphs)
    table_count = len(tables)

    return {
        "file_type": "docx",
        "status": "valid",
        "metadata": metadata,
        
        # Quantity statistics
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "has_text_layer": has_text,
        
        # Technical decisions for Step 2 (Processing)
        "requires_table_extraction": table_count > 0,
        "requires_text_processing": has_text,
        "is_empty": not (has_text or table_count > 0),
        
        # Additional info for logs
        "word_version": "Office Open XML (OOXML)"
    }